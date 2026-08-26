from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(
    "/Users/joantisdale/Documents/Codex/2026-08-24/"
    "referenced-chatgpt-conversation-this-is-an/outputs/"
    "MeteoLabX_explicaciones_parametros_convectivos.docx"
)
OUTPUT = Path(__file__).with_name("MeteoLabX_explicaciones_parametros_convectivos_EBWD.docx")


def set_language(run):
    r_pr = run._element.get_or_add_rPr()
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = r_pr.makeelement(qn("w:lang"), {})
        r_pr.append(lang)
    lang.set(qn("w:val"), "es-ES")


def add_rich_paragraph(anchor, parts, style="Normal"):
    paragraph = anchor.insert_paragraph_before(style=style)
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        set_language(run)
    return paragraph


def add_plain_paragraph(anchor, text, style="Normal"):
    return add_rich_paragraph(anchor, [(text, False, False)], style=style)


def add_heading(anchor, text, level):
    paragraph = add_plain_paragraph(anchor, text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_step(anchor, label, text):
    return add_rich_paragraph(
        anchor,
        [(label, True, False), (text, False, False)],
        style="List Number",
    )


def add_equation(anchor, pieces):
    paragraph = anchor.insert_paragraph_before(style="Ecuación")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    for text, formatting in pieces:
        run = paragraph.add_run(text)
        run.font.name = "Cambria Math"
        r_pr = run._element.get_or_add_rPr()
        r_pr.get_or_add_rFonts().set(qn("w:ascii"), "Cambria Math")
        r_pr.get_or_add_rFonts().set(qn("w:hAnsi"), "Cambria Math")
        run.font.size = Pt(12)
        run.bold = formatting.get("bold", False)
        run.italic = formatting.get("italic", False)
        run.font.subscript = formatting.get("subscript", False)
        run.font.superscript = formatting.get("superscript", False)
        set_language(run)
    return paragraph


def add_ecape_equation(anchor, label, parcel_subscript):
    add_equation(
        anchor,
        [
            (label + " = ∫", {}),
            ("LFC", {"subscript": True}),
            ("EL", {"superscript": True}),
            ("  g  [", {}),
            ("T", {"italic": True}),
            (parcel_subscript, {"subscript": True}),
            ("(entr)", {"superscript": True}),
            (" − ", {}),
            ("T", {"italic": True}),
            ("v,e", {"subscript": True}),
            ("] / ", {}),
            ("T", {"italic": True}),
            ("v,e", {"subscript": True}),
            ("  d", {}),
            ("z", {"italic": True}),
        ],
    )


def set_numbering_group(document, paragraphs):
    style_num_id = document.styles["List Number"]._element.pPr.numPr.numId.val
    numbering = document.part.numbering_part.element
    source_num = next(
        item
        for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_num_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max(
        int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))
    ) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)

    for paragraph in paragraphs:
        num_pr = paragraph._element.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = next_num_id


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def main():
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    start = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.startswith("EBWD ·")
    )
    end = next(
        i
        for i, paragraph in enumerate(document.paragraphs[start + 1 :], start + 1)
        if paragraph.text.strip() == "Próximos parámetros"
    )
    for paragraph in list(document.paragraphs[start:end]):
        remove_paragraph(paragraph)

    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Próximos parámetros"
    )

    add_heading(anchor, "EBWD · Diferencia vectorial efectiva del viento", 1)

    add_heading(anchor, "Qué mide", 2)
    add_rich_paragraph(
        anchor,
        [
            ("La ", False, False),
            ("EBWD", True, False),
            (
                " (Effective Bulk Wind Difference) mide cuánto cambia el viento, como vector, "
                "dentro de la mitad inferior de la profundidad efectiva de la tormenta. El cálculo "
                "empieza en la base de la capa de entrada que realmente puede alimentar la convección "
                "y termina a mitad de camino hacia el nivel de equilibrio de la parcela más inestable.",
                False,
                False,
            ),
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            (
                "Por eso no usa siempre una capa fija desde superficie hasta 6 km. Frente a CIZ6, "
                "EBWD se adapta tanto a la altura del inflow como a la profundidad de la tormenta. "
                "La diferencia es especialmente útil en convección elevada y en tormentas más "
                "someras o más profundas de lo habitual.",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Cálculo paso a paso", 2)
    ebwd_steps = []
    ebwd_steps.append(add_step(
        anchor,
        "Construir el perfil. ",
        "MeteoLabX combina la superficie con los niveles isobáricos de AROME para obtener temperatura, humedad, altura y componentes U/V en una misma columna.",
    ))
    ebwd_steps.append(add_step(
        anchor,
        "Encontrar la base efectiva. ",
        "Se prueban parcelas desde la superficie hasta 500 hPa. La primera que cumple simultáneamente los dos criterios siguientes fija la base de la capa:",
    ))
    add_equation(
        anchor,
        [
            ("CAPE", {}),
            ("parcela", {"subscript": True}),
            (" ≥ 100 J kg", {}),
            ("−1", {"superscript": True}),
            (",     CIN", {}),
            ("parcela", {"subscript": True}),
            (" ≥ −250 J kg", {}),
            ("−1", {"superscript": True}),
        ],
    )
    ebwd_steps.append(add_step(
        anchor,
        "Situar el techo efectivo. ",
        "Se toma el punto medio, en altura, entre esa base y el nivel de equilibrio (EL) de la parcela más inestable (MU):",
    ))
    add_equation(
        anchor,
        [
            ("z", {"italic": True}),
            ("top", {"subscript": True}),
            (" = ", {}),
            ("z", {"italic": True}),
            ("base", {"subscript": True}),
            (" + ½ (", {}),
            ("z", {"italic": True}),
            ("EL,MU", {"subscript": True}),
            (" − ", {}),
            ("z", {"italic": True}),
            ("base", {"subscript": True}),
            (")", {}),
        ],
    )
    ebwd_steps.append(add_step(
        anchor,
        "Calcular la diferencia vectorial. ",
        "U y V se interpolan linealmente en la base y en el techo. Después se resta el vector de la base al del techo y se calcula su módulo:",
    ))
    add_equation(
        anchor,
        [
            ("EBWD = | ", {}),
            ("V", {"bold": True, "italic": True}),
            ("(", {}),
            ("z", {"italic": True}),
            ("top", {"subscript": True}),
            (") − ", {}),
            ("V", {"bold": True, "italic": True}),
            ("(", {}),
            ("z", {"italic": True}),
            ("base", {"subscript": True}),
            (") |", {}),
        ],
    )

    add_heading(anchor, "Cómo interpretar el mapa", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "Los colores expresan la magnitud de la EBWD: cuanto mayor es el valor, mayor es el "
                "cambio del viento dentro de la capa efectiva. En presencia de inestabilidad y un "
                "mecanismo de disparo, una EBWD mayor suele favorecer tormentas más organizadas y "
                "persistentes porque ayuda a separar la corriente ascendente de la descendente y de "
                "la precipitación.",
                False,
                False,
            )
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            (
                "Como referencia operativa, el entorno se vuelve progresivamente más favorable para "
                "supercélulas cuando la EBWD entra aproximadamente en el intervalo de ",
                False,
                False,
            ),
            ("25 a 40 kt", True, False),
            (
                " o lo supera. No es un umbral aislado: la respuesta depende también de CAPE, CIN, "
                "la hodógrafa, la profundidad de la tormenta y el movimiento de las células.",
                False,
                False,
            ),
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            ("Las flechas muestran la ", False, False),
            ("orientación del vector diferencia", True, False),
            (
                ". No representan el viento de un nivel concreto ni el desplazamiento de las tormentas. "
                "Para estudiar rotación en niveles bajos deben compararse con ESRH, CIZ1, el movimiento "
                "de tormenta y la forma completa de la hodógrafa.",
                False,
                False,
            ),
        ],
    )

    add_heading(anchor, "Límites", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "EBWD describe la organización potencial de la convección profunda, pero no es por sí "
                "sola un indicador de tornados, granizo o líneas convectivas. Tampoco incorpora el "
                "movimiento de tormenta, la curvatura de la hodógrafa, el cold pool ni las interacciones "
                "entre células. Si ninguna parcela cumple los criterios de CAPE y CIN, la capa efectiva "
                "no existe y EBWD queda sin definir.",
                False,
                False,
            )
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            (
                "MeteoLabX calcula el diagnóstico directamente a partir del perfil termodinámico y U/V "
                "de AROME, siguiendo la definición de capa efectiva de Thompson et al. (2007); no llama "
                "a SHARPpy para obtener EBWD. El resultado es sensible a la resolución vertical y a la "
                "interpolación entre niveles.",
                False,
                False,
            )
        ],
    )
    note = add_plain_paragraph(anchor, "Diagnóstico MeteoLabX · perfil termodinámico y U/V AROME")
    note.runs[0].italic = True

    add_heading(anchor, "MU-ECAPE · CAPE con arrastre de la parcela más inestable", 1)

    add_heading(anchor, "Qué representa", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "MU-ECAPE es la CAPE nativa de AROME asociada a la parcela más inestable de las "
                "capas bajas. El producto del modelo incluye sus propios efectos de dilución o "
                "arrastre, es decir, la mezcla de aire ambiental con la parcela mientras asciende. "
                "MeteoLabX usa el nombre MU-ECAPE para distinguir este campo de su MUCAPE convencional, "
                "calculada sin arrastre.",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Interpretación", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "El mapa busca la parte del entorno con mayor flotabilidad potencial. Por eso puede "
                "mostrar inestabilidad elevada aunque la superficie sea estable. Un valor alto indica "
                "que, incluso después de la dilución representada por AROME, queda una cantidad importante "
                "de energía disponible para acelerar una corriente ascendente.",
                False,
                False,
            )
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            (
                "La comparación más útil es con MUCAPE MLX y con ML-ECAPE. Si MU-ECAPE supera claramente "
                "a ML-ECAPE, la capa más inestable puede estar elevada o ser poco representativa del promedio "
                "de la capa baja. Si MU-ECAPE es mucho menor que MUCAPE MLX, el producto de AROME está "
                "representando una reducción importante de la flotabilidad por dilución. Estas diferencias "
                "orientan la lectura, pero no permiten deducir por sí solas el esquema exacto del modelo.",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Cálculo", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "MeteoLabX descarga el campo nativo CONVECTIVE_AVAILABLE_POTENTIAL_ENERGY publicado por "
                "AROME, conserva su valor en J/kg y lo representa. No reconstruye la parcela, no recalcula "
                "la energía y no llama a SHARPpy. La forma física general de una CAPE con parcela diluida es:",
                False,
                False,
            )
        ],
    )
    add_ecape_equation(anchor, "ECAPE", "v,p")
    mu_steps = []
    mu_steps.append(add_step(
        anchor,
        "Origen del dato. ",
        "La selección de la parcela MU, su trayectoria y el arrastre pertenecen al producto nativo de AROME.",
    ))
    mu_steps.append(add_step(
        anchor,
        "Tratamiento MeteoLabX. ",
        "El valor se descarga y se muestra sin correcciones ni combinación con el Lifted Index de MLX, que se calcula sin arrastre.",
    ))

    add_heading(anchor, "Notas y limitaciones", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "La API pública no documenta la formulación exacta de la temperatura virtual de la "
                "parcela diluida, la tasa de arrastre, el cierre ni el procedimiento preciso de selección "
                "de la parcela. Por ello, MU-ECAPE no debe compararse uno a uno con MUCAPE MLX como si solo "
                "cambiara una constante conocida. Como toda CAPE, tampoco informa por sí sola sobre el "
                "disparo, la inhibición, la cizalladura o el modo convectivo.",
                False,
                False,
            )
        ],
    )
    note = add_plain_paragraph(anchor, "AROME · parcela MU con arrastre")
    note.runs[0].italic = True

    add_heading(anchor, "ML-ECAPE · CAPE con arrastre de capa mezclada", 1)

    add_heading(anchor, "Qué representa", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "ML-ECAPE es la CAPE nativa de AROME para una parcela representativa de una capa baja "
                "mezclada, con la dilución o arrastre incluidos por el propio producto del modelo. MeteoLabX "
                "la etiqueta así para diferenciarla de MLCAPE MLX, que usa una parcela de capa mezclada "
                "convencional sin arrastre.",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Interpretación", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "Al representar un promedio de la capa baja, ML-ECAPE suele ser menos sensible que una "
                "parcela superficial a máximos muy locales de temperatura o humedad. Describe mejor el "
                "reservorio medio de flotabilidad disponible para convección enraizada en la capa límite.",
                False,
                False,
            )
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            (
                "Debe leerse junto con MU-ECAPE. Valores parecidos sugieren que la capa baja mezclada "
                "representa bien la parcela más favorable; una MU-ECAPE claramente mayor puede señalar "
                "una capa elevada más inestable o una franja especialmente cálida y húmeda que el promedio "
                "ML suaviza. La comparación con MLCAPE MLX ayuda a apreciar la reducción asociada al "
                "producto con arrastre.",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Cálculo", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "MeteoLabX descarga el campo nativo MEAN_LAYER_CAPE publicado por AROME y lo muestra en "
                "J/kg sin reconstruir la parcela ni modificar la energía. Su forma física general puede "
                "expresarse como:",
                False,
                False,
            )
        ],
    )
    add_ecape_equation(anchor, "ML-ECAPE", "v,p,ML")
    ml_steps = []
    ml_steps.append(add_step(
        anchor,
        "Origen del dato. ",
        "La selección de la capa ML, la profundidad de mezcla y el arrastre son internos del campo AROME.",
    ))
    ml_steps.append(add_step(
        anchor,
        "Tratamiento MeteoLabX. ",
        "El campo se descarga y se representa sin correcciones posteriores.",
    ))

    add_heading(anchor, "Notas y limitaciones", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "La API pública no documenta la profundidad exacta de mezcla ni el esquema de arrastre, "
                "por lo que no debe suponerse que ML-ECAPE usa exactamente los 100 hPa inferiores empleados "
                "por la MLCAPE convencional de MeteoLabX. El campo expresa energía potencial; para valorar "
                "si esa energía puede realizarse y qué tipo de tormenta podría producir, debe combinarse "
                "con CIN, forzamiento, humedad, cizalladura y estructura vertical.",
                False,
                False,
            )
        ],
    )
    note = add_plain_paragraph(anchor, "AROME · parcela ML con arrastre")
    note.runs[0].italic = True

    set_numbering_group(document, ebwd_steps)
    set_numbering_group(document, mu_steps)
    set_numbering_group(document, ml_steps)
    anchor.paragraph_format.page_break_before = False
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
