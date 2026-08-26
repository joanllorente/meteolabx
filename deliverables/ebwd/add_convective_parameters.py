from pathlib import Path
import shutil

from docx import Document

from rewrite_ebwd import (
    add_equation,
    add_heading,
    add_plain_paragraph,
    add_rich_paragraph,
    add_step,
    remove_paragraph,
    set_numbering_group,
)


SOURCE = Path(__file__).with_name(
    "MeteoLabX_explicaciones_parametros_convectivos_EBWD.docx"
)
OUTPUT = Path(__file__).with_name(
    "MeteoLabX_explicaciones_parametros_convectivos_completo.docx"
)


def add_note(anchor, text):
    paragraph = add_plain_paragraph(anchor, text)
    paragraph.runs[0].italic = True


def add_simple_equation(anchor, text):
    return add_equation(anchor, [(text, {})])


def add_mucape(anchor):
    add_heading(anchor, "MUCAPE + MULI · Parcela más inestable", 1)

    add_heading(anchor, "Qué representa", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "MUCAPE es la CAPE convencional, sin arrastre, de la parcela más inestable de los "
                "300 hPa inferiores del perfil. MeteoLabX la representa mediante colores y dibuja en "
                "isolíneas el MULI, el Lifted Index calculado para exactamente la misma parcela MU.",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Interpretación", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "MUCAPE identifica la mayor flotabilidad potencial, aunque la parcela de origen esté "
                "elevada. MULI describe la flotabilidad de esa parcela a 500 hPa: un valor negativo "
                "significa que la parcela llega más cálida que el ambiente. MUCAPE alta y MULI muy "
                "negativo refuerzan la señal de inestabilidad, pero no garantizan que exista disparo ni "
                "que la parcela pueda superar la inhibición.",
                False,
                False,
            )
        ],
    )
    add_rich_paragraph(
        anchor,
        [
            (
                "Como el ascenso no incluye arrastre, carga de agua ni mezcla lateral, MUCAPE funciona "
                "como un límite superior idealizado de la energía de la corriente ascendente. Conviene "
                "compararla con MU-ECAPE para apreciar la reducción representada por AROME y con CIN "
                "para valorar si la parcela puede alcanzar el nivel de libre convección (LFC).",
                False,
                False,
            )
        ],
    )

    add_heading(anchor, "Cálculo", 2)
    add_rich_paragraph(
        anchor,
        [
            (
                "MeteoLabX aplica su algoritmo vectorizado de parcela. Selecciona el máximo de temperatura "
                "potencial equivalente de Bolton entre la superficie y 300 hPa por encima, eleva la parcela "
                "en seco hasta el LCL y pseudoadiabáticamente después. La flotabilidad usa temperatura "
                "virtual y la energía se integra por trapecios.",
                False,
                False,
            )
        ],
    )
    add_simple_equation(anchor, "p_MU = arg max[p_s−300 ≤ p ≤ p_s]  θ_e(p)")
    add_simple_equation(
        anchor,
        "MUCAPE = ∫[LFC, EL] g (T_v,p − T_v,e) / T_v,e  dz",
    )
    add_simple_equation(
        anchor,
        "MULI = T_e(500 hPa) − T_p,MU(500 hPa)",
    )
    steps = [
        add_step(anchor, "Seleccionar la parcela MU. ", "Se busca el máximo θe dentro de los 300 hPa inferiores."),
        add_step(anchor, "Elevar la parcela. ", "El ascenso es seco hasta el LCL y pseudoadiabático por encima, con el condensado eliminado y sin entrainment."),
        add_step(anchor, "Integrar la CAPE. ", "Se suman en altura las capas con flotabilidad virtual positiva; las alturas proceden de integración hipsométrica."),
        add_step(anchor, "Calcular MULI. ", "Se resta la temperatura de la parcela MU a la ambiental en 500 hPa."),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "El diagnóstico no llama a params.cape de SHARPpy. CAPE y LI describen flotabilidad, no el disparo, el modo convectivo ni la intensidad final de la tormenta.",
    )
    add_note(anchor, "Diagnóstico MeteoLabX · MUCAPE + isolíneas MULI")
    return steps


def add_mlcape(anchor):
    add_heading(anchor, "MLCAPE + MLLI · Parcela de capa mezclada", 1)

    add_heading(anchor, "Qué representa", 2)
    add_plain_paragraph(
        anchor,
        "MLCAPE es la CAPE sin arrastre de una parcela representativa de los 100 hPa inferiores. MeteoLabX mezcla esa capa mediante la temperatura potencial y la razón de mezcla medias. Los colores muestran MLCAPE y las isolíneas muestran el MLLI de la misma parcela.",
    )

    add_heading(anchor, "Interpretación", 2)
    add_plain_paragraph(
        anchor,
        "MLCAPE suaviza picos muy locales de temperatura o humedad y suele representar mejor una capa límite bien mezclada. Es apropiada para convección que ingiere un espesor de aire, no solo las condiciones del primer nivel cercano a 2 m.",
    )
    add_plain_paragraph(
        anchor,
        "MLLI negativo refuerza la señal de inestabilidad a 500 hPa. Una SBCAPE mucho mayor que MLCAPE puede revelar una capa superficial cálida o húmeda extremadamente fina; una MUCAPE claramente mayor que MLCAPE puede indicar que la capa más inestable está elevada.",
    )

    add_heading(anchor, "Cálculo", 2)
    add_plain_paragraph(
        anchor,
        "MeteoLabX promedia θ y la razón de mezcla r en los 100 hPa inferiores, reconstruye temperatura y punto de rocío a la presión de superficie y eleva la parcela con el mismo esquema pseudoadiabático, virtual y sin arrastre utilizado para MU.",
    )
    add_simple_equation(
        anchor,
        "θ̄ = (1/Δp) ∫[p_s−100, p_s] θ dp,     r̄ = (1/Δp) ∫[p_s−100, p_s] r dp",
    )
    add_simple_equation(
        anchor,
        "MLCAPE = ∫[LFC, EL] B dz,     MLLI = T_e(500) − T_p,ML(500)",
    )
    steps = [
        add_step(anchor, "Definir la capa ML100. ", "Se toman los 100 hPa situados inmediatamente sobre la presión de superficie."),
        add_step(anchor, "Mezclar sus propiedades. ", "Se promedian θ y r y se reconstruyen T y Td de la parcela a p_s."),
        add_step(anchor, "Elevar e integrar. ", "La parcela asciende seca al LCL y pseudoadiabáticamente al EL; la CAPE usa flotabilidad virtual."),
        add_step(anchor, "Calcular MLLI. ", "Las isolíneas usan la temperatura a 500 hPa de exactamente la misma parcela ML."),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "El cálculo no incluye arrastre ni carga de condensado y no llama a params.cape de SHARPpy. Debe combinarse con CIN, forzamiento y cizalladura.",
    )
    add_note(anchor, "Diagnóstico MeteoLabX · MLCAPE + isolíneas MLLI")
    return steps


def add_sbcape(anchor):
    add_heading(anchor, "SBCAPE + SBLI · Parcela superficial", 1)

    add_heading(anchor, "Qué representa", 2)
    add_plain_paragraph(
        anchor,
        "SBCAPE es la CAPE sin arrastre de una parcela que parte de las condiciones de superficie del perfil, construidas con temperatura y punto de rocío cercanos a 2 m. Los colores muestran SBCAPE y las isolíneas el SBLI de esa misma parcela.",
    )

    add_heading(anchor, "Interpretación", 2)
    add_plain_paragraph(
        anchor,
        "Es especialmente sensible al ciclo diurno, las brisas, los frentes de racha y las piscinas frías. Resulta útil para convección claramente enraizada en superficie, aunque puede exagerar una capa cálida o húmeda demasiado delgada.",
    )
    add_plain_paragraph(
        anchor,
        "SBCAPE elevada con SBLI negativo indica flotabilidad potencial de una parcela superficial, pero no asegura que venza la inhibición. Si SBCAPE disminuye mientras MUCAPE permanece alta, la inestabilidad puede haberse elevado por encima de una capa superficial estable.",
    )

    add_heading(anchor, "Cálculo", 2)
    add_plain_paragraph(
        anchor,
        "MeteoLabX inserta T y Td de superficie como primer nivel y eleva esa parcela con el mismo esquema vectorizado pseudoadiabático, de temperatura virtual y sin arrastre usado en las demás CAPE MLX.",
    )
    add_simple_equation(
        anchor,
        "SBCAPE = ∫[LFC, EL] g (T_v,p,SFC − T_v,e) / T_v,e  dz",
    )
    add_simple_equation(
        anchor,
        "SBLI = T_e(500 hPa) − T_p,SFC(500 hPa)",
    )
    steps = [
        add_step(anchor, "Fijar el origen. ", "Se usan la presión de superficie y T/Td del primer nivel."),
        add_step(anchor, "Elevar la parcela. ", "El ascenso es seco hasta el LCL y pseudoadiabático por encima."),
        add_step(anchor, "Calcular energía e índice. ", "La CAPE se integra en altura y el LI se evalúa en 500 hPa."),
        add_step(anchor, "Representar el resultado. ", "SBCAPE se muestra en colores y SBLI en isolíneas."),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "No se incluyen mezcla lateral, carga de condensado ni forzamiento dinámico. El cálculo no llama a params.cape de SHARPpy y debe interpretarse junto con CIN y el espesor real de la capa superficial favorable.",
    )
    add_note(anchor, "Diagnóstico MeteoLabX · SBCAPE + isolíneas SBLI")
    return steps


def add_dcape(anchor):
    add_heading(anchor, "DCAPE · Energía potencial de corrientes descendentes", 1)

    add_heading(anchor, "Qué representa", 2)
    add_plain_paragraph(
        anchor,
        "DCAPE (Downdraft CAPE) estima la energía potencial que una parcela saturada de niveles medios podría convertir en aceleración descendente al bajar pseudoadiabáticamente hasta la superficie.",
    )

    add_heading(anchor, "Interpretación", 2)
    add_plain_paragraph(
        anchor,
        "Valores altos suelen aparecer con aire seco en niveles medios y una capa baja con fuerte gradiente térmico, condiciones favorables al enfriamiento evaporativo y a descensos intensos. DCAPE ayuda a detectar entornos de downburst, pero no predice la racha exacta que alcanzará el suelo.",
    )
    add_plain_paragraph(
        anchor,
        "La realización depende de la precipitación disponible, la carga de hidrometeoros, la profundidad de evaporación, la mezcla y la organización de la tormenta. Debe combinarse con reflectividad o precipitación, humedad, racha máxima y estructura convectiva.",
    )

    add_heading(anchor, "Cálculo", 2)
    add_plain_paragraph(
        anchor,
        "MeteoLabX reproduce el procedimiento SPC/SHARPpy. Dentro de los 400 hPa inferiores busca la capa móvil de 100 hPa con menor θe media; la parcela parte del centro de esa capa, se satura a la temperatura de bulbo húmedo y desciende húmedo-adiabáticamente.",
    )
    add_simple_equation(
        anchor,
        "p_0 = p_base,min(θ̄_e) − 50 hPa",
    )
    add_simple_equation(
        anchor,
        "DCAPE = −R_d ∫[p_s, p_0] (T_e − T_p) d ln p",
    )
    steps = [
        add_step(anchor, "Buscar capas candidatas. ", "Sus bases coinciden con niveles nativos entre p_s y p_s−400 hPa; θe se muestrea cada 5 hPa dentro de cada capa de 100 hPa."),
        add_step(anchor, "Elegir el origen. ", "La parcela parte del centro de la capa de 100 hPa con menor θe media."),
        add_step(anchor, "Saturar la parcela. ", "La temperatura inicial se lleva a bulbo húmedo."),
        add_step(anchor, "Descender e integrar. ", "Se usa wetlift y una integral trapezoidal con temperatura ordinaria, sin corrección de temperatura virtual, igual que params.dcape."),
        add_step(anchor, "Aplicar el fallback. ", "Si SHARPpy no está disponible, se usa la inversión vectorizada de θe saturada."),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "DCAPE expresa energía potencial de descenso. No incorpora por sí sola cantidad de precipitación, carga de agua, organización ni transferencia completa del momento al suelo.",
    )
    add_note(anchor, "Diagnóstico MeteoLabX · perfil termodinámico AROME")
    return steps


def add_cell_motion(anchor):
    add_heading(anchor, "Movimiento de células ordinarias", 1)

    add_heading(anchor, "Qué representa", 2)
    add_plain_paragraph(
        anchor,
        "Es una estimación de la componente advectiva del movimiento de una célula convectiva ordinaria. Se calcula con el viento medio ponderado por presión dentro de la nube de una parcela ML100, entre su LCL y su EL.",
    )

    add_heading(anchor, "Interpretación", 2)
    add_plain_paragraph(
        anchor,
        "Los colores muestran la velocidad estimada y las líneas de corriente la dirección de traslación por el flujo medio. Sirve para anticipar hacia dónde se desplazaría una célula no supercelular y cuánto tiempo podría permanecer sobre una zona.",
    )
    add_plain_paragraph(
        anchor,
        "No incluye propagación por nuevos desarrollos, cold pools, splitting de supercélulas, interacción con fronteras ni anclaje orográfico. El movimiento de un MCS puede ser muy distinto porque combina advección y propagación.",
    )

    add_heading(anchor, "Cálculo", 2)
    add_plain_paragraph(
        anchor,
        "MeteoLabX calcula primero el LCL y el EL de la parcela ML100. Después integra U y V por trapecios sobre presión, recorta cada capa isobárica al intervalo nuboso y divide por la profundidad de presión.",
    )
    add_simple_equation(
        anchor,
        "C⃗_cel = [1 / (p_LCL − p_EL)] ∫[p_EL, p_LCL] V⃗(p) dp",
    )
    add_simple_equation(anchor, "C_cel = √(C_u² + C_v²)")
    steps = [
        add_step(anchor, "Definir la nube. ", "LCL y EL proceden de la parcela ML100 calculada por MeteoLabX."),
        add_step(anchor, "Construir el perfil de viento. ", "U/V proceden de la superficie y los niveles isobáricos de AROME."),
        add_step(anchor, "Promediar por presión. ", "Se calcula el solape exacto de cada capa con el intervalo LCL–EL y se integra linealmente por trapecios."),
        add_step(anchor, "Representar el vector. ", "La magnitud aparece en colores y la dirección mediante streamlines."),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "No es un movimiento Bunkers ni Corfidi y no representa supercélulas o sistemas propagativos. Describe únicamente la advección por el viento medio dentro de la nube ML100.",
    )
    add_note(anchor, "Diagnóstico MeteoLabX · viento medio ML100 LCL–EL")
    return steps


def add_ship(anchor):
    add_heading(anchor, "SHIP · Significant Hail Parameter", 1)

    add_heading(anchor, "Qué representa", 2)
    add_plain_paragraph(
        anchor,
        "SHIP es un índice compuesto adimensional para identificar entornos favorables al granizo significativo. Combina MUCAPE, humedad de la parcela MU, gradiente térmico 700–500 hPa, temperatura a 500 hPa, cizalladura 0–6 km y altura de congelación.",
    )

    add_heading(anchor, "Interpretación", 2)
    add_plain_paragraph(
        anchor,
        "Valores crecientes indican un entorno progresivamente más favorable porque combinan ascensos intensos, humedad, crecimiento en aire frío y ventilación por cizalladura. SHIP no es tamaño de granizo, probabilidad ni garantía de granizo en una celda concreta.",
    )
    add_plain_paragraph(
        anchor,
        "Los umbrales proceden del contexto operativo del SPC de Estados Unidos y no están calibrados para la Península Ibérica. Debe evaluarse junto con modo convectivo, cota de congelación, CAPE en la zona de crecimiento del granizo y observaciones.",
    )

    add_heading(anchor, "Cálculo", 2)
    add_plain_paragraph(
        anchor,
        "Cuando SHARPpy está instalado, MeteoLabX llama celda a celda a sharppy.sharptab.params.ship. Le entrega la MUCAPE MLX, la razón de mezcla MU, el lapse rate 700–500, T500, la cizalladura geométrica superficie–6 km y la altura AGL de 0 °C. El fallback reproduce la misma función y sus limitadores.",
    )
    add_simple_equation(
        anchor,
        "SHIP_0 = −(MUCAPE · r_MU · Γ_700–500 · T_500 · BWD_0–6) / 42 000 000",
    )
    add_simple_equation(
        anchor,
        "SHIP = max(0, SHIP_0) · f_C · f_Γ · f_F",
    )
    add_simple_equation(
        anchor,
        "f_C=min(1,MUCAPE/1300),   f_Γ=min(1,Γ_700–500/5.8),   f_F=min(1,z_0°C/2400)",
    )
    steps = [
        add_step(anchor, "Limitar humedad. ", "La razón de mezcla MU se restringe al intervalo 11–13,6 g/kg."),
        add_step(anchor, "Limitar cizalladura y temperatura. ", "BWD0–6 se restringe a 7–27 m/s y T500 a un máximo de −5,5 °C."),
        add_step(anchor, "Calcular los campos verticales. ", "Γ700–500 usa T y altura hipsométrica AROME; la altura de 0 °C es AGL."),
        add_step(anchor, "Aplicar reductores. ", "CAPE baja, lapse rate débil o nivel de congelación bajo reducen el resultado mediante f_C, f_Γ y f_F."),
        add_step(anchor, "Truncar la salida. ", "El resultado final es adimensional y no puede ser negativo."),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "La MUCAPE usada es el diagnóstico MLX sin arrastre. Esta formulación emplea BWD geométrica superficie–6 km y no EBWD. SHIP resume ingredientes ambientales; no sustituye el análisis del modo convectivo ni de la microfísica real de cada tormenta.",
    )
    add_note(anchor, "Diagnóstico MeteoLabX · formulación SPC sobre perfiles AROME")
    return steps


def remove_completed_placeholders(document):
    paragraphs = list(document.paragraphs)
    start = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "Próximos parámetros")
    for label in ("CAPE", "SHIP"):
        index = next(
            i
            for i, paragraph in enumerate(paragraphs[start + 1 :], start + 1)
            if paragraph.text.strip() == label
        )
        remove_paragraph(paragraphs[index])
        if index + 1 < len(paragraphs) and paragraphs[index + 1].text.strip() == "Pendiente de desarrollar.":
            remove_paragraph(paragraphs[index + 1])


def main():
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Próximos parámetros"
    )

    numbering_groups = [
        add_mucape(anchor),
        add_mlcape(anchor),
        add_sbcape(anchor),
        add_dcape(anchor),
        add_cell_motion(anchor),
        add_ship(anchor),
    ]
    for group in numbering_groups:
        set_numbering_group(document, group)

    remove_completed_placeholders(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
