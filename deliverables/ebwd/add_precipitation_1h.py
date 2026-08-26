from pathlib import Path
import shutil

from docx import Document

from rewrite_ebwd import (
    add_equation,
    add_heading,
    add_plain_paragraph,
    add_step,
    set_numbering_group,
)


SOURCE = Path(__file__).with_name(
    "MeteoLabX_explicaciones_parametros_convectivos_completo.docx"
)
OUTPUT = Path(__file__).with_name(
    "MeteoLabX_explicaciones_parametros_meteorologicos.docx"
)


def replace_text_preserving_runs(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def add_precipitation_1h(anchor):
    add_heading(anchor, "Precipitación en 1 hora", 1)

    add_heading(anchor, "Qué representa", 2)
    add_plain_paragraph(
        anchor,
        "Es la precipitación total acumulada durante la hora inmediatamente anterior a la hora válida del mapa. Incluye lluvia y todas las demás fases de precipitación expresadas como equivalente de agua líquida.",
    )

    add_heading(anchor, "Interpretación", 2)
    add_plain_paragraph(
        anchor,
        "El mapa muestra la cantidad acumulada en una hora, no la severidad de la tormenta ni su intensidad instantánea. El total depende tanto de cuánto precipita el sistema como del tiempo que permanece sobre cada punto.",
    )
    add_plain_paragraph(
        anchor,
        "Por ello, una célula relativamente débil pero lenta o estacionaria puede dejar acumulaciones horarias altas. En cambio, una tormenta muy intensa o severa que se desplaza rápidamente puede producir acumulaciones pequeñas en cada punto, aunque genere granizo, viento fuerte o actividad eléctrica intensa.",
    )
    add_plain_paragraph(
        anchor,
        "La precipitación convectiva tiene un error importante de fase y posición. Una celda prevista pocos kilómetros fuera de su lugar puede producir un error local grande aunque el patrón meteorológico general sea correcto. El mapa no debe interpretarse como una medición puntual exacta.",
    )

    add_heading(anchor, "Cálculo", 2)
    add_plain_paragraph(
        anchor,
        "Es un campo WCS nativo de AROME: TOTAL_PRECIPITATION con sufijo temporal PT1H. MeteoLabX selecciona la hora solicitada, limita a cero cualquier valor negativo y aplica la equivalencia de agua 1 kg/m² = 1 mm.",
    )
    add_equation(
        anchor,
        [
            ("P", {"italic": True}),
            ("1h", {"subscript": True}),
            ("(", {}),
            ("t", {"italic": True}),
            (") = ∫", {}),
            ("t−1h", {"subscript": True}),
            ("t", {"superscript": True}),
            ("  R(τ) dτ", {}),
        ],
    )
    add_equation(
        anchor,
        [
            ("1 kg m", {}),
            ("−2", {"superscript": True}),
            (" = 1 mm", {}),
        ],
    )
    steps = [
        add_step(
            anchor,
            "Seleccionar el campo nativo. ",
            "La cobertura es TOTAL_PRECIPITATION__GROUND_OR_WATER_SURFACE_PT1H.",
        ),
        add_step(
            anchor,
            "Asignar el intervalo. ",
            "El mapa válido a la hora t representa exclusivamente el intervalo (t−1 h, t].",
        ),
        add_step(
            anchor,
            "Mostrar la acumulación. ",
            "Los valores se expresan en milímetros y no se suman con las horas vecinas en este producto.",
        ),
    ]

    add_heading(anchor, "Notas y limitaciones", 2)
    add_plain_paragraph(
        anchor,
        "Este campo no diagnostica severidad convectiva: no informa directamente de granizo, rachas, rayos, rotación ni organización. Tampoco distingue la fase que alcanzó el suelo, porque todas se convierten a equivalente líquido. Para estudiar la duración completa del episodio deben consultarse varias horas o el acumulado desde el inicio del RUN.",
    )
    note = add_plain_paragraph(anchor, "AROME · TOTAL_PRECIPITATION · superficie · PT1H")
    note.runs[0].italic = True
    return steps


def main():
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    title = next(
        paragraph
        for paragraph in document.paragraphs
        if "MeteoLabX · Explicaciones de parámetros" in paragraph.text
    )
    replace_text_preserving_runs(title, "parámetros convectivos", "parámetros meteorológicos")

    objective = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Objetivo.")
    )
    replace_text_preserving_runs(objective, "productos convectivos", "productos meteorológicos")

    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Próximos parámetros"
    )
    steps = add_precipitation_1h(anchor)
    set_numbering_group(document, steps)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
